#!/usr/bin/env python3
"""
File Loader Module for Resume Parser POC

This module provides functionality to load and extract text content from various file formats
including PDF and DOCX files. It normalizes the extracted text for further processing.

Author: Resume Parser POC
Version: 1.0
"""

import os
import re
from typing import Optional


class FileLoader:
    """
    A utility class for loading and extracting text from different file formats.
    
    Supported formats:
    - PDF files (using PyMuPDF/fitz)
    - DOCX files (using python-docx)
    - TXT files (plain text)
    """
    
    @staticmethod
    def load_file(file_path: str) -> str:
        """
        Load and extract text from a file based on its extension.
        
        Args:
            file_path (str): Path to the file to be loaded
            
        Returns:
            str: Extracted and normalized text content
            
        Raises:
            FileNotFoundError: If the file doesn't exist
            ValueError: If the file format is not supported
            Exception: For other file processing errors
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        # Get file extension
        _, ext = os.path.splitext(file_path.lower())
        
        try:
            if ext == '.pdf':
                return FileLoader.load_pdf(file_path)
            elif ext == '.docx':
                return FileLoader.load_docx(file_path)
            elif ext == '.txt':
                return FileLoader.load_txt(file_path)
            else:
                raise ValueError(f"Unsupported file format: {ext}")
        except Exception as e:
            raise Exception(f"Error processing file {file_path}: {str(e)}")
    
    @staticmethod
    def load_pdf(file_path: str) -> str:
        """
        Extract text content from a PDF file using PyMuPDF.
        
        Args:
            file_path (str): Path to the PDF file
            
        Returns:
            str: Extracted text content
        """
        try:
            import fitz  # PyMuPDF
        except ImportError:
            raise ImportError("PyMuPDF (fitz) is required for PDF processing. Install with: pip install PyMuPDF")
        
        text = ""
        try:
            # Open the PDF document
            doc = fitz.open(file_path)
            
            # Extract text from each page
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                text += page.get_text()
                text += "\n"  # Add page separator
            
            doc.close()
            
        except Exception as e:
            raise Exception(f"Error reading PDF file: {str(e)}")
        
        return FileLoader.normalize_text(text)
    
    @staticmethod
    def load_docx(file_path: str) -> str:
        """
        Extract text content from a DOCX file using python-docx.
        
        Args:
            file_path (str): Path to the DOCX file
            
        Returns:
            str: Extracted text content
        """
        try:
            from docx import Document
        except ImportError:
            raise ImportError("python-docx is required for DOCX processing. Install with: pip install python-docx")
        
        try:
            # Load the document
            doc = Document(file_path)
            
            # Extract text from paragraphs
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
        except Exception as e:
            raise Exception(f"Error reading DOCX file: {str(e)}")
        
        return FileLoader.normalize_text(text)
    
    @staticmethod
    def load_txt(file_path: str) -> str:
        """
        Load content from a plain text file.
        
        Args:
            file_path (str): Path to the text file
            
        Returns:
            str: File content
        """
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
                text = file.read()
        except Exception as e:
            raise Exception(f"Error reading text file: {str(e)}")
        
        return FileLoader.normalize_text(text)
    
    @staticmethod
    def normalize_text(text: str) -> str:
        """
        Normalize extracted text by cleaning and standardizing format.
        
        Args:
            text (str): Raw extracted text
            
        Returns:
            str: Normalized text
        """
        if not text:
            return ""
        
        # Remove control characters except newlines and tabs
        text = re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F\x7F-\x9F]', '', text)
        
        # Normalize excessive whitespace within lines (but preserve line breaks)
        lines = text.split('\n')
        normalized_lines = []
        for line in lines:
            # Remove excessive whitespace within each line
            normalized_line = re.sub(r'[ \t]+', ' ', line).strip()
            normalized_lines.append(normalized_line)
        
        # Rejoin with single newlines
        text = '\n'.join(normalized_lines)
        
        # Remove excessive empty lines (more than 2 consecutive)
        text = re.sub(r'\n\s*\n\s*\n+', '\n\n', text)
        
        # Remove leading/trailing whitespace
        text = text.strip()
        
        return text
    
    @staticmethod
    def get_file_info(file_path: str) -> dict:
        """
        Get metadata information about a file.
        
        Args:
            file_path (str): Path to the file
            
        Returns:
            dict: File metadata including size, extension, etc.
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        stat = os.stat(file_path)
        _, ext = os.path.splitext(file_path.lower())
        
        return {
            'path': file_path,
            'filename': os.path.basename(file_path),
            'extension': ext,
            'size_bytes': stat.st_size,
            'size_kb': round(stat.st_size / 1024, 2),
            'is_supported': ext in ['.pdf', '.docx', '.txt']
        }


# Example usage and testing
if __name__ == "__main__":
    # Example usage
    loader = FileLoader()
    
    # Test with a sample file (replace with actual file path)
    try:
        file_path = "sample_resume.pdf"
        if os.path.exists(file_path):
            text = loader.load_file(file_path)
            print(f"Extracted text length: {len(text)} characters")
            print(f"First 200 characters: {text[:200]}...")
        else:
            print(f"Sample file {file_path} not found")
    except Exception as e:
        print(f"Error: {e}")